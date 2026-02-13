# -*- coding: utf-8 -*-
"""从目录加载多种格式的企业文档（PDF、Word、TXT、Markdown）。"""
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from logger_config import logger


def _load_txt_or_md(path: Path) -> list[Document]:
    """加载 .txt / .md：直接按 utf-8 读入，避免 Windows 路径/默认编码问题。"""
    try:
        text = path.read_text(encoding="utf-8")
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"source": str(path), "filename": path.name})]
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="gbk")
            if not text.strip():
                return []
            return [Document(page_content=text, metadata={"source": str(path), "filename": path.name})]
        except Exception as e:
            logger.error(f"加载文件失败（编码）{path}: {e}")
            return []
    except Exception as e:
        logger.error(f"加载文件失败 {path}: {e}")
        return []


def _load_docx(path: Path) -> list[Document]:
    """使用 python-docx 加载 .docx 文件（仅支持 .docx，不支持旧版 .doc）。"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(parts)
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"source": str(path), "filename": path.name})]
    except Exception as e:
        logger.warning(f"python-docx 加载失败 {path}: {e}")
        return []


def _load_doc(path: Path) -> list[Document]:
    """加载旧版 .doc 文件：尝试使用 unstructured 库（需要额外依赖）。"""
    try:
        from unstructured.partition.doc import partition_doc
        
        elements = partition_doc(str(path))
        # elements 是 Element 对象列表，提取文本内容
        text_parts = []
        for el in elements:
            if hasattr(el, 'text') and el.text:
                text_parts.append(el.text)
            elif hasattr(el, '__str__'):
                text_parts.append(str(el))
        text = "\n\n".join(text_parts)
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"source": str(path), "filename": path.name})]
    except ImportError:
        logger.warning(
            f"无法加载 .doc 文件 {path.name}：需要安装 unstructured 库。"
            f"请运行: pip install 'unstructured[local]'"
            f"或将该文件转换为 .docx 格式。"
        )
        return []
    except Exception as e:
        logger.warning(f"unstructured 加载 .doc 失败 {path}: {e}，建议转换为 .docx")
        return []


# 支持的后缀与加载方式（Loader 类或 (path -> list[Document]) 函数）
LOADER_MAP = {
    ".pdf": PyPDFLoader,
    ".txt": _load_txt_or_md,
    ".md": _load_txt_or_md,
    ".markdown": _load_txt_or_md,
}
# docx/doc 使用自定义函数
try:
    import docx  # noqa: F401
    LOADER_MAP[".docx"] = _load_docx
except ImportError:
    pass
# .doc（旧版 Word）单独处理，使用 unstructured
LOADER_MAP[".doc"] = _load_doc


def _load_file(path: Path) -> list[Document]:
    """加载单个文件，返回 Document 列表。"""
    suffix = path.suffix.lower()
    loader = LOADER_MAP.get(suffix)
    if not loader:
        logger.warning(f"暂不支持的文件格式: {path.name} ({suffix})")
        return []
    try:
        if callable(loader) and not hasattr(loader, "load"):
            return loader(path)
        loader_instance = loader(str(path))
        docs = loader_instance.load()
        for d in docs:
            d.metadata.setdefault("source", str(path))
            d.metadata.setdefault("filename", path.name)
        return docs
    except Exception as e:
        logger.error(f"加载文件失败 {path}: {e}")
        return []


def load_documents_from_directory(directory: str | Path) -> list[Document]:
    """
    从指定目录递归加载所有支持格式的文档。
    支持: .pdf, .txt, .md, .docx 等。
    """
    directory = Path(directory)
    if not directory.is_dir():
        logger.warning(f"目录不存在或不是目录: {directory}")
        return []

    all_docs: list[Document] = []
    supported = set(LOADER_MAP.keys())
    for path in directory.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in supported:
            continue
        docs = _load_file(path)
        all_docs.extend(docs)
        if docs:
            logger.info(f"已加载: {path.name} -> {len(docs)} 个片段")

    logger.info(f"共加载 {len(all_docs)} 个文档块，来自目录 {directory}")
    return all_docs
